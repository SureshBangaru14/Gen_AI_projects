import os
import tempfile

from pathlib import Path

from pdf2image import convert_from_path

import pytesseract

from docx import Document


class JDProcessor:

    # ========================================================
    # CONSTRUCTOR
    # ========================================================

    def __init__(
        self,
        input_method=None,
        pdf_file=None,
        docx_file=None,
        pasted_text=None
    ):

        self.input_method = input_method

        self.pdf_file = pdf_file

        self.docx_file = docx_file

        self.pasted_text = pasted_text


    # ========================================================
    # MAIN PROCESS
    # ========================================================

    def process(self):

        # ====================================================
        # PDF
        # ====================================================

        if self.input_method == "Upload PDF":

            return self.process_pdf()


        # ====================================================
        # DOCX
        # ====================================================

        if self.input_method == "Upload DOCX":

            return self.process_docx()


        # ====================================================
        # PASTE TEXT
        # ====================================================

        if self.input_method == "Paste Text":

            return self.process_pasted_text()


        raise ValueError(

            "Invalid Job Description input method. "

            "Expected: Upload PDF, Upload DOCX, "
            "or Paste Text."

        )


    # ========================================================
    # PDF PROCESSING
    # ========================================================

    def process_pdf(self):

        if not self.pdf_file:

            raise ValueError(
                "Job Description PDF was not provided."
            )


        # ====================================================
        # STREAMLIT UPLOADED FILE
        # ====================================================

        file_bytes = (
            self.pdf_file.getvalue()
        )


        original_file_name = (
            self.pdf_file.name
        )


        temp_pdf_path = None


        try:

            # =================================================
            # CREATE TEMP PDF
            # =================================================

            with tempfile.NamedTemporaryFile(

                delete=False,

                suffix=".pdf"

            ) as temp_file:

                temp_file.write(
                    file_bytes
                )


                temp_pdf_path = (
                    temp_file.name
                )


            # =================================================
            # PDF → IMAGES
            # =================================================

            images = (
                convert_from_path(

                    temp_pdf_path,

                    dpi=300

                )
            )


            page_texts = []


            # =================================================
            # OCR EACH PAGE
            # =================================================

            for page_number, image in enumerate(

                images,

                start=1

            ):

                text = (
                    pytesseract.image_to_string(
                        image
                    )
                )


                text = (
                    self.clean_text(
                        text
                    )
                )


                if text:

                    page_texts.append(
                        text
                    )


            # =================================================
            # PAGE 1 + PAGE 2 + PAGE 3
            # =================================================

            full_text = (

                "\n\n".join(
                    page_texts
                )

            )


            # =================================================
            # FINAL CLEAN
            # =================================================

            full_text = (
                self.clean_text(
                    full_text
                )
            )


            if not full_text:

                raise ValueError(

                    f"No text could be extracted "
                    f"from JD PDF: "
                    f"{original_file_name}"

                )


            return full_text


        finally:

            # =================================================
            # DELETE TEMP FILE
            # =================================================

            if (

                temp_pdf_path

                and

                os.path.exists(
                    temp_pdf_path
                )

            ):

                os.remove(
                    temp_pdf_path
                )


    # ========================================================
    # DOCX PROCESSING
    # ========================================================

    def process_docx(self):

        if not self.docx_file:

            raise ValueError(
                "Job Description DOCX was not provided."
            )


        original_file_name = (
            self.docx_file.name
        )


        file_bytes = (
            self.docx_file.getvalue()
        )


        temp_docx_path = None


        try:

            # =================================================
            # CREATE TEMP DOCX
            # =================================================

            with tempfile.NamedTemporaryFile(

                delete=False,

                suffix=".docx"

            ) as temp_file:

                temp_file.write(
                    file_bytes
                )


                temp_docx_path = (
                    temp_file.name
                )


            # =================================================
            # READ DOCX
            # =================================================

            document = Document(
                temp_docx_path
            )


            text_parts = []


            # =================================================
            # PARAGRAPHS
            # =================================================

            for paragraph in (
                document.paragraphs
            ):

                text = (
                    paragraph.text.strip()
                )


                if text:

                    text_parts.append(
                        text
                    )


            # =================================================
            # TABLES
            # =================================================

            for table in (
                document.tables
            ):

                for row in table.rows:

                    row_values = []


                    for cell in row.cells:

                        cell_text = (
                            cell.text.strip()
                        )


                        if cell_text:

                            row_values.append(
                                cell_text
                            )


                    if row_values:

                        text_parts.append(

                            " | ".join(
                                row_values
                            )

                        )


            # =================================================
            # COMBINE
            # =================================================

            full_text = (

                "\n\n".join(
                    text_parts
                )

            )


            # =================================================
            # CLEAN
            # =================================================

            full_text = (
                self.clean_text(
                    full_text
                )
            )


            if not full_text:

                raise ValueError(

                    f"No text could be extracted "
                    f"from JD DOCX: "
                    f"{original_file_name}"

                )


            return full_text


        finally:

            # =================================================
            # DELETE TEMP FILE
            # =================================================

            if (

                temp_docx_path

                and

                os.path.exists(
                    temp_docx_path
                )

            ):

                os.remove(
                    temp_docx_path
                )


    # ========================================================
    # PASTED TEXT
    # ========================================================

    def process_pasted_text(self):

        if self.pasted_text is None:

            raise ValueError(
                "Job Description text was not provided."
            )


        text = str(
            self.pasted_text
        )


        text = (
            self.clean_text(
                text
            )
        )


        if not text:

            raise ValueError(

                "Job Description text "
                "cannot be empty."

            )


        return text


    # ========================================================
    # CLEAN TEXT
    # ========================================================

    def clean_text(
        self,
        text
    ):

        if not text:

            return ""


        # ====================================================
        # NORMALIZE LINE BREAKS
        # ====================================================

        text = text.replace(
            "\r\n",
            "\n"
        )


        text = text.replace(
            "\r",
            "\n"
        )


        # ====================================================
        # REMOVE TRAILING SPACES
        # ====================================================

        lines = []


        for line in text.split(
            "\n"
        ):

            line = line.strip()


            if line:

                lines.append(
                    line
                )


        text = "\n".join(
            lines
        )


        # ====================================================
        # NORMALIZE MULTIPLE SPACES
        # ====================================================

        while "  " in text:

            text = text.replace(
                "  ",
                " "
            )


        # ====================================================
        # NORMALIZE MULTIPLE NEWLINES
        # ====================================================

        while "\n\n\n" in text:

            text = text.replace(
                "\n\n\n",
                "\n\n"
            )


        return text.strip()